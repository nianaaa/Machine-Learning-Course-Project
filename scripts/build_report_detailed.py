from __future__ import annotations

import json
import math
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORK_DIR = Path("/mnt/sdc/zoujunjie/mlearn_power_coursework")
OUT_DOCX = WORK_DIR / "reports" / "mlearn_power_detailed_report.docx"

MODEL_NAMES = {
    "lstm": "LSTM",
    "transformer": "Transformer",
    "pvg_itransformer": "PVG-iTransformer",
    "pvg_no_time": "PVG w/o Time Patch",
    "pvg_no_variable": "PVG w/o Variable",
    "pvg_no_gate": "PVG w/o Gate",
}

CORE_MODELS = ["lstm", "transformer", "pvg_itransformer"]
ABLATION_REPORT_MODELS = [
    "pvg_itransformer",
    "pvg_no_time",
    "pvg_no_variable",
    "pvg_no_gate",
]


def east_asia(run, font: str = "宋体") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def paragraph(doc: Document, text: str, first_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(22) if first_indent else Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    east_asia(run)
    run.font.size = Pt(11)


def small_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    east_asia(run)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("555555")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    east_asia(run)
    run.bold = bold
    run.font.size = Pt(9.2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    caption: str | None = None,
    left_cols: set[int] | None = None,
) -> None:
    left_cols = left_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in left_cols else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, align=align)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    if caption:
        add_caption(doc, caption)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    east_asia(run)
    run.font.size = Pt(9)


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(10)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.2)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def fmt(mean: float, std: float) -> str:
    return f"{mean:.2f} ± {std:.2f}"


def model_name(name: str) -> str:
    return MODEL_NAMES.get(name, name)


def metric(summary: pd.DataFrame, model: str, horizon: int, col: str) -> float:
    row = summary[(summary["model"] == model) & (summary["horizon"] == horizon)].iloc[0]
    return float(row[col])


def compare_pct(new: float, base: float) -> str:
    pct = (new - base) / base * 100.0
    return f"低 {abs(pct):.2f}%" if pct <= 0 else f"高 {pct:.2f}%"


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("家庭电力消耗时间序列预测实验报告")
    east_asia(run, "黑体")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("专硕机器学习课程考核 - LSTM、Transformer 与 PVG-iTransformer")
    east_asia(run)
    run.font.size = Pt(11)


def main() -> None:
    report_dir = WORK_DIR / "reports"
    report_dir.mkdir(parents=True, exist_ok=True)

    summary = pd.read_csv(WORK_DIR / "results" / "metrics_summary.csv")
    core_summary = summary[summary["model"].isin(CORE_MODELS)].copy()
    ablation_summary = summary[summary["model"].isin(ABLATION_REPORT_MODELS)].copy()
    runs = pd.read_csv(WORK_DIR / "results" / "metrics_runs.csv")
    meta = json.loads((WORK_DIR / "results" / "run_metadata.json").read_text(encoding="utf-8"))
    impute = pd.read_csv(WORK_DIR / "data" / "processed" / "minute_imputation_summary.csv")
    station = pd.read_csv(WORK_DIR / "data" / "processed" / "weather_station_suresnes_summary.csv")
    daily = pd.read_csv(WORK_DIR / "data" / "processed" / "daily_power.csv", parse_dates=["date"])

    data_meta = meta["data"]
    h90_meta = meta["horizon_90"]
    h365_meta = meta["horizon_365"]
    feature_dim = len(h90_meta["feature_cols"])
    weather_cols = data_meta["weather_cols"]
    weather_text = "、".join(weather_cols)
    station_name = data_meta["weather_station_name"]
    station_distance = float(data_meta["weather_station_distance_km"])

    split_idx = h90_meta["split_idx"]
    train_start = daily["date"].iloc[0].date()
    train_end = daily["date"].iloc[split_idx - 1].date()
    test_start = daily["date"].iloc[split_idx].date()
    test_end = daily["date"].iloc[-1].date()

    missing_minutes = int(impute["missing_minutes"].max())
    min_candidates = int(impute["min_candidates"].min())
    max_candidates = int(impute["max_candidates"].max())
    mean_candidates = float(impute["mean_candidates"].mean())

    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("摘要", level=1)
    paragraph(
        doc,
        "本文按照《2026 年专硕机器学习课程项目》要求，完成家庭电力消耗多变量时间序列预测实验。"
        "实验使用 UCI Individual household electric power consumption 原始分钟级数据，并结合 data.gouv.fr "
        "提供的 Meteo-France 92 省月度气象数据。预测目标为按天汇总后的 daily global_active_power，输入为过去 "
        "90 天的多变量历史序列，输出分别为未来 90 天短期预测曲线和未来 365 天长期预测曲线。"
        "本文实现并比较三类主模型：LSTM、标准 Transformer，以及受 PatchTST 和 iTransformer 启发设计的 "
        "PVG-iTransformer。三类主模型在 90 天和 365 天任务上分别独立训练，每种设置使用 5 个随机种子重复实验，"
        "并报告 MSE 与 MAE 的均值和标准差；此外补充 PVG-iTransformer 的三组消融实验，以检验时间 Patch 分支、"
        "变量分支和门控融合模块的贡献。",
    )
    paragraph(
        doc,
        "数据处理方面，本文先补齐完整 1 分钟时间索引，再对分钟级缺失值使用同月份内 t±7、t±14、t±21、t±28 天的"
        "同星期几、同分钟时刻观测值均值进行填补，然后按 PDF 要求聚合到日尺度。天气变量选用 SURESNES 气象站的 "
        f"{weather_text} 四个完整月度变量；PDF 中列出的 NBJBROU 因该站 48 个月中缺失 47 个月，未纳入主实验。"
        "实验采用 65/35 时间划分，并在测试阶段使用 rolling-origin evaluation。主模型比较显示，PVG-iTransformer "
        "在 90 天和 365 天预测上均优于 LSTM 与标准 Transformer。消融实验进一步显示，变量分支贡献最明显，移除变量分支会"
        "显著增大误差；移除时间 Patch 分支后误差反而下降，说明当前小样本日度任务中变量 token 视角更稳定，时间 Patch 分支可能引入额外噪声。",
    )
    small_note(doc, "关键词：家庭电力消耗；多变量时间序列预测；LSTM；Transformer；PVG-iTransformer；rolling-origin evaluation")

    doc.add_heading("1. 问题介绍", level=1)
    doc.add_heading("1.1 课程任务与实际意义", level=2)
    paragraph(
        doc,
        "家庭电力消耗预测是智能家居、能耗管理和电网调度中的典型应用。家庭用电受季节、工作日与周末行为、"
        "住户活动、用电设备、电压电流状态以及气象条件等多种因素影响，因此具有明显的多变量时间序列特征。"
        "准确预测未来每天的总有功功率，有助于居民理解自身用电行为，也可以为峰谷用电管理、异常用电识别和"
        "电网负荷调度提供参考。",
    )
    paragraph(
        doc,
        "课程 PDF 将任务定义为：基于过去 90 天的数据曲线，分别预测未来 90 天和未来 365 天的总有功功率变化曲线。"
        "其中未来 90 天属于短期预测，未来 365 天属于长期预测；两种预测长度需要分别训练，长期预测模型的参数不能"
        "用于短期预测。方法上分为三部分：第一部分使用 LSTM，第二部分使用 Transformer，第三部分使用自己提出的"
        "改进模型，且第三部分以原理的新颖程度为首要评价标准，性能为次要评价标准。",
    )
    add_table(
        doc,
        ["PDF 要求", "本文实现方式"],
        [
            ["报告结构", "按问题介绍、模型、结果与分析、讨论四部分组织，并补充摘要和参考文献"],
            ["预测任务", "过去 90 天输入，分别预测未来 90 天和未来 365 天 daily global_active_power"],
            ["三种方法", "LSTM、Transformer、PVG-iTransformer 三类主模型完整实现和比较"],
            ["消融实验", "补充移除时间 Patch 分支、变量分支和门控模块的三组 PVG 变体实验"],
            ["独立训练", "90 天与 365 天分别构造数据窗口并独立训练模型，参数不共享"],
            ["评价指标", "使用 MSE 与 MAE，在原始 daily sum 尺度计算误差"],
            ["重复实验", "每个模型、每个 horizon 使用 seeds 42-46 共 5 轮实验，报告均值和 std"],
            ["结果截图", "报告中贴入结果表截图和 power/Ground Truth 曲线对比图"],
            ["参考文献", "列出 UCI、Meteo-France、LSTM、Transformer、PatchTST、iTransformer 等文献"],
            ["代码提交", "代码目录已整理，正式提交前需将目录上传 GitHub 并替换报告中的链接占位"],
        ],
        [1.65, 4.85],
        "表 1 PDF 要求与本文实现对应关系",
        left_cols={0, 1},
    )

    doc.add_heading("1.2 数据来源与变量说明", level=2)
    paragraph(
        doc,
        "原始电力数据来自 UCI Machine Learning Repository 的 Individual household electric power consumption 数据集，"
        "记录法国一户家庭从 2006 年 12 月到 2010 年 11 月的分钟级用电信息。原始字段包括 global_active_power、"
        "global_reactive_power、voltage、global_intensity、sub_metering_1、sub_metering_2 和 sub_metering_3。"
        "根据 PDF 提示，还可以由总有功功率和三个分表能耗计算剩余能耗 sub_metering_remainder。",
    )
    paragraph(
        doc,
        "天气数据来自 data.gouv.fr 的 Meteo-France 月度基础气候数据。考虑到 UCI 数据为法国 Hauts-de-Seine 区域家庭，"
        "本文下载 92 省月度文件，并在站点层面检查缺失情况。最终选择 SURESNES 站点，编号 NUM_POSTE=92073001，"
        f"经纬度约为 {station.iloc[0]['LAT']:.6f}, {station.iloc[0]['LON']:.6f}，与家庭所在地近似距离约 {station_distance:.1f} km。"
        "该站在实验所需的 48 个月内 RR、NBJRR1、NBJRR5、NBJRR10 均无缺失。RR 原始单位为毫米的十分之一，按 PDF 说明除以 10 后使用。",
    )

    doc.add_heading("1.3 数据划分与可复现环境", level=2)
    paragraph(
        doc,
        "课程 PDF 提到数据集主要分为 train 和 test 两部分，具体见 train.csv 和 tes.csv。由于远端目录中未发现课程额外提供的"
        " train.csv 与 tes.csv，本文从 PDF 指定的公开 UCI 数据源自动下载原始数据，并按时间顺序构造 65/35 的训练/测试划分。"
        f"最终日度数据共 {len(daily)} 天，时间范围为 {daily['date'].iloc[0].date()} 至 {daily['date'].iloc[-1].date()}；"
        f"训练段为 {train_start} 至 {train_end}，测试段为 {test_start} 至 {test_end}。",
    )
    add_table(
        doc,
        ["项目", "数值或说明"],
        [
            ["工作目录", str(WORK_DIR)],
            ["Python 环境", "/mnt/sdc/zoujunjie/miniconda3/envs/mlearn"],
            ["PyTorch 版本", meta["device"]["torch"]],
            ["训练设备", meta["device"]["name"]],
            ["日度样本数", str(len(daily))],
            ["训练/测试比例", "65% / 35%，按时间顺序切分"],
            ["训练日期", f"{train_start} 至 {train_end}"],
            ["测试日期", f"{test_start} 至 {test_end}"],
        ],
        [1.7, 4.8],
        "表 2 数据划分与实验环境",
        left_cols={1},
    )

    doc.add_heading("2. 模型", level=1)
    doc.add_heading("2.1 数据处理与监督样本构造", level=2)
    paragraph(
        doc,
        "PDF 提示分钟级数据需要自行处理。若直接把含缺失的分钟数据按天求和，缺失分钟会系统性压低日总电量；因此本文先补齐完整"
        " 1 分钟时间索引，再在分钟级填补缺失。对于任意缺失分钟 t，本文在同一月份内查找 t±7、t±14、t±21、t±28 天的"
        "同星期几、同分钟时刻观测值，并对可用候选取均值。这个规则既利用了家庭用电的周周期，又避免跨月引入明显不同的季节背景。"
        f"经统计，7 个原始电力变量各有 {missing_minutes} 个缺失分钟；所有缺失点至少有 {min_candidates} 个候选，最多有 "
        f"{max_candidates} 个候选，平均候选数为 {mean_candidates:.2f}。",
    )
    paragraph(
        doc,
        "完成分钟级填补后，本文按 PDF 要求进行日度聚合：global_active_power、global_reactive_power、sub_metering_1、"
        "sub_metering_2 和 sub_metering_3 按天求和；voltage 与 global_intensity 按天求平均；sub_metering_remainder "
        "由 (global_active_power * 1000 / 60) - (sub_metering_1 + sub_metering_2 + sub_metering_3) 在分钟级计算后按天求和。"
        "天气变量为月度数据，因此将同一月份的 RR、NBJRR1、NBJRR5、NBJRR10 复制到当月每天。由于这些变量并非真实日度天气，"
        "报告中明确将其称为月度气象背景变量。",
    )
    add_table(
        doc,
        ["特征类别", "字段", "日度处理方式"],
        [
            ["目标与功率", "global_active_power, global_reactive_power", "按天求和"],
            ["电力状态", "voltage, global_intensity", "按天求平均"],
            ["分表能耗", "sub_metering_1/2/3", "按天求和"],
            ["派生能耗", "sub_metering_remainder", "由总有功功率和三个分表计算后按天求和"],
            ["天气变量", weather_text, "SURESNES 月度天气值复制到当月每天，RR 除以 10"],
            ["日期周期", "dayofweek_sin/cos, month_sin/cos", "正余弦编码"],
        ],
        [1.2, 2.6, 2.7],
        "表 3 特征构造与日度聚合方式",
        left_cols={1, 2},
    )
    paragraph(
        doc,
        "标准化参数只在训练段日序列上计算，避免测试集信息泄露。输入特征使用训练段均值和标准差标准化，输出目标使用训练段 "
        "global_active_power 的均值和标准差标准化。模型训练时在标准化空间优化 MSE，测试时将预测值反标准化回原始 daily sum "
        "尺度，再计算 MSE 和 MAE。",
    )
    paragraph(
        doc,
        "监督样本采用滑动窗口构造。输入窗口长度固定为 90 天，输出窗口长度 H 分别为 90 和 365。训练窗口要求输出区间完全位于"
        "训练段内，测试窗口要求预测起点位于测试段内。测试阶段采用 rolling-origin evaluation：模型参数固定后，在测试时间段内"
        "滚动构造多个测试窗口，并汇总所有测试窗口、所有预测点的误差。最终输入维度为 "
        f"{feature_dim}，其中包括 8 个电力/分表相关变量、{len(weather_cols)} 个天气变量和 4 个日期周期变量。"
        f"90 天任务得到 {h90_meta['train_windows']} 个训练窗口和 {h90_meta['test_windows']} 个测试窗口；"
        f"365 天任务得到 {h365_meta['train_windows']} 个训练窗口和 {h365_meta['test_windows']} 个测试窗口。",
    )
    add_formula(doc, "MSE = (1/N) Σ_i (y_i - ŷ_i)^2")
    add_formula(doc, "MAE = (1/N) Σ_i |y_i - ŷ_i|")
    paragraph(
        doc,
        "其中 N 表示所有 rolling-origin 测试窗口中预测点的总数。MSE 对大误差更敏感，MAE 更直观地表示平均每天的绝对偏差。"
        "由于本任务的目标是按天求和后的 global_active_power，MSE 的绝对值会比分钟级或 kWh 尺度更大，解释结果时需要同时参考"
        " MAE 和 RMSE。",
    )

    doc.add_heading("2.2 LSTM 基线模型", level=2)
    paragraph(
        doc,
        "LSTM 是处理序列依赖的经典循环神经网络。本文使用单层 LSTM 对过去 90 天的多变量输入进行编码，取最后一个时间步的隐藏状态"
        "作为历史窗口的压缩表示，再经过 LayerNorm、Dropout 和线性预测头直接输出未来 H 天的完整曲线。90 天和 365 天任务分别"
        "训练独立模型，输出维度分别为 90 和 365。",
    )
    paragraph(
        doc,
        f"LSTM 的输入张量形状为 [B, 90, {feature_dim}]。模型按时间顺序递归更新隐藏状态，因此较适合捕捉相邻日期之间的局部连续性。"
        "本文采用 direct multi-output 方式一次性输出完整预测曲线，而不是逐日递推，避免递推式多步预测中的误差逐步累积。",
    )

    doc.add_heading("2.3 Transformer 基线模型", level=2)
    paragraph(
        doc,
        "标准 Transformer 将每天的多变量观测作为一个时间 token。具体做法是先将每个日度特征向量映射到 64 维隐藏空间，加入"
        "正弦位置编码后送入两层 Transformer Encoder。自注意力机制能够直接建模任意两个日期之间的关系，因此理论上比循环结构"
        "更适合捕捉较长时间范围内的模式变化。模型使用最后一个时间 token 的编码表示，通过线性预测头输出未来 H 天曲线。",
    )
    paragraph(
        doc,
        "Transformer 的优点是并行计算和长依赖建模能力较强；缺点是在日度样本数量有限时，参数量和注意力结构可能更容易受训练样本"
        "规模、正则化强度和随机初始化影响。因此本文用 5 个随机种子重复实验，并报告标准差来衡量稳定性。",
    )

    doc.add_heading("2.4 PVG-iTransformer 改进模型", level=2)
    paragraph(
        doc,
        "第三部分提出 PVG-iTransformer，即 Patch-Variable Gated iTransformer，中文可称为基于时间片段与变量反转门控融合的 "
        "Transformer。该模型受 PatchTST 的时间 patch 思想和 iTransformer 的变量反转思想启发，但针对本任务设计为双分支结构。"
        "它不是简单堆叠更多层数，而是改变 Transformer 的 token 构造方式，使模型同时从时间片段和变量关系两个视角理解家庭用电序列。",
    )
    paragraph(
        doc,
        "时间 Patch 分支将 90 天输入切分为长度 7、步长 3 的周级片段，并强制包含最后一个 patch，最终得到 29 个 patch token。"
        "每个 patch 展平后经线性层映射到 64 维隐藏空间，加入可学习位置嵌入，再送入两层 Transformer Encoder，最后通过 mean pooling "
        "得到时间视角表示 f_time。该分支强调近期周周期、阶段变化和局部趋势。",
    )
    paragraph(
        doc,
        "变量分支将输入转置为 [B, D, 90]，把每个变量过去 90 天的历史作为一个 variable token。每个变量 token 经 Linear(90, 64) "
        "映射后，加入变量嵌入和变量组嵌入。变量组包括目标功率、电力上下文、分表能耗、天气变量和日历变量。该分支借鉴 iTransformer "
        "的反转思想，重点学习变量之间的依赖关系，并取 Global_active_power 对应 token 作为变量视角表示 f_var。",
    )
    paragraph(
        doc,
        "双分支输出通过门控融合模块组合：g = sigmoid(W[f_time; f_var] + b)，f = g * f_time + (1 - g) * f_var。"
        "当某个样本更依赖近期时间趋势时，门控可以提高 f_time 的权重；当跨变量关系更重要时，门控可以提高 f_var 的权重。"
        "融合表示最后经过 MLP 直接输出未来 90 或 365 天预测值。",
    )
    add_table(
        doc,
        ["模块", "输入 token", "作用"],
        [
            ["时间 Patch 分支", "29 个周级 patch token", "学习局部周期、阶段变化和近期趋势"],
            ["变量分支", f"{feature_dim} 个 variable token", "学习目标功率、电力状态、分表、天气和日历变量关系"],
            ["门控融合", "f_time 与 f_var", "根据样本自适应平衡时间模式和变量关系"],
            ["预测头", "融合表示 f", "直接输出未来 90 或 365 天总有功功率"],
        ],
        [1.35, 2.05, 3.1],
        "表 4 PVG-iTransformer 结构组成",
        left_cols={2},
    )
    paragraph(
        doc,
        "为验证各组成部分的实际贡献，本文在完整 PVG-iTransformer 之外补充三组消融变体。PVG w/o Time Patch 移除时间 Patch "
        "分支，仅保留变量分支和预测头；PVG w/o Variable 移除变量分支，仅保留时间 Patch 分支和预测头；PVG w/o Gate 保留两个"
        "分支，但把可学习门控融合替换为 f = 0.5 * f_time + 0.5 * f_var 的固定平均融合。三组变体使用与完整模型完全相同的数据划分、"
        "标准化、优化器、训练轮数、随机种子和 rolling-origin 评价流程，因此结果差异可以主要解释为结构组件差异。",
    )
    paragraph(doc, "PVG-iTransformer 的前向过程可概括为如下伪代码，其中 H 表示预测长度，H=90 或 H=365：")
    add_code_block(
        doc,
        [
            f"Input: X in R^{{B x 90 x D}}, D={feature_dim}",
            "Patch tokens: split X into weekly patches, Linear(patch) -> P",
            "Time branch: f_time = MeanPool(TransformerEncoder(P + position_embedding))",
            "Variable tokens: transpose X to R^{B x D x 90}, Linear(history) -> V",
            "Variable branch: f_var = target_token(TransformerEncoder(V + var_embedding + group_embedding))",
            "Gate: g = sigmoid(W concat(f_time, f_var) + b)",
            "Fusion: f = g * f_time + (1 - g) * f_var",
            "Output: y_hat = MLP(f) in R^{B x H}",
        ],
    )

    doc.add_heading("2.5 训练设置", level=2)
    paragraph(
        doc,
        "三类主模型及 PVG 消融变体均使用 AdamW 优化器，学习率 1e-3，weight decay 为 1e-4，batch size 为 32，训练 30 个 epoch。"
        "损失函数为标准化目标空间中的 MSE。每个 epoch 后记录训练损失，最终保留训练损失最低的模型状态。随机种子为 "
        "42、43、44、45、46，共 5 轮。所有模型在同一数据划分、同一标准化方式、同一训练轮数和同一评价流程下比较。",
    )

    doc.add_heading("3. 结果与分析", level=1)
    result_rows = []
    for _, row in core_summary.iterrows():
        result_rows.append(
            [
                model_name(row["model"]),
                str(int(row["horizon"])),
                fmt(row["mse_mean"], row["mse_std"]),
                fmt(row["mae_mean"], row["mae_std"]),
                str(int(row["runs"])),
            ]
        )
    add_table(
        doc,
        ["模型", "预测天数", "MSE 均值±std", "MAE 均值±std", "轮数"],
        result_rows,
        [1.15, 1.0, 1.65, 1.65, 0.7],
        "表 5 三种主模型五轮实验平均结果",
    )
    doc.add_picture(str(WORK_DIR / "figures" / "metrics_summary_table.png"), width=Inches(5.8))
    add_caption(doc, "图 1 主实验与消融实验结果汇总截图")

    pvg90_mse = metric(summary, "pvg_itransformer", 90, "mse_mean")
    pvg90_mae = metric(summary, "pvg_itransformer", 90, "mae_mean")
    tr90_mse = metric(summary, "transformer", 90, "mse_mean")
    tr90_mae = metric(summary, "transformer", 90, "mae_mean")
    lstm90_mse = metric(summary, "lstm", 90, "mse_mean")
    lstm90_mae = metric(summary, "lstm", 90, "mae_mean")
    pvg365_mse = metric(summary, "pvg_itransformer", 365, "mse_mean")
    pvg365_mae = metric(summary, "pvg_itransformer", 365, "mae_mean")
    tr365_mse = metric(summary, "transformer", 365, "mse_mean")
    tr365_mae = metric(summary, "transformer", 365, "mae_mean")
    lstm365_mse = metric(summary, "lstm", 365, "mse_mean")
    lstm365_mae = metric(summary, "lstm", 365, "mae_mean")

    paragraph(
        doc,
        f"从表 5 可以看出，90 天短期预测中，PVG-iTransformer 的 MSE 为 {pvg90_mse:.2f}，MAE 为 {pvg90_mae:.2f}，"
        f"均为三种主模型最低。相对于标准 Transformer，PVG 的 MSE {compare_pct(pvg90_mse, tr90_mse)}、MAE "
        f"{compare_pct(pvg90_mae, tr90_mae)}；相对于 LSTM，PVG 的 MSE {compare_pct(pvg90_mse, lstm90_mse)}、MAE "
        f"{compare_pct(pvg90_mae, lstm90_mae)}。这说明短期任务中，周级 patch 对近期趋势的刻画和变量分支对跨变量关系的建模"
        "共同降低了平均误差。",
    )
    paragraph(
        doc,
        f"365 天长期预测中，PVG-iTransformer 的 MSE 为 {pvg365_mse:.2f}，MAE 为 {pvg365_mae:.2f}，同样为三种主模型最低。"
        f"相对于 LSTM，PVG 的 MSE {compare_pct(pvg365_mse, lstm365_mse)}、MAE {compare_pct(pvg365_mae, lstm365_mae)}；"
        f"相对于标准 Transformer，PVG 的 MSE {compare_pct(pvg365_mse, tr365_mse)}、MAE {compare_pct(pvg365_mae, tr365_mae)}。"
        "长期预测需要一次性输出 365 个未来点，训练窗口比 90 天任务更少，因此整体更难；PVG 在该任务上仍优于两个基线，"
        "说明其双视角 token 构造对长期季节轮廓也有帮助。",
    )
    rmse_min = math.sqrt(float(summary["mse_mean"].min()))
    rmse_max = math.sqrt(float(summary["mse_mean"].max()))
    paragraph(
        doc,
        "需要注意，课程要求对 global_active_power 按天求和，因此 MSE 在日累积尺度上计算，数值看起来较大。"
        f"这些 MSE 对应的 RMSE 约为 {rmse_min:.2f} 至 {rmse_max:.2f}；若按 kWh 近似折算，即除以 60，约为 "
        f"{rmse_min / 60.0:.2f} 至 {rmse_max / 60.0:.2f} kWh。因此评价时不能只看 MSE 绝对值，也应同时看 MAE、RMSE "
        "以及不同模型之间的相对差异。",
    )

    doc.add_page_break()
    doc.add_heading("3.1 PVG-iTransformer 消融实验", level=2)
    ablation_rows = []
    for _, row in ablation_summary.iterrows():
        full_mse = metric(summary, "pvg_itransformer", int(row["horizon"]), "mse_mean")
        mse_delta = "基准" if row["model"] == "pvg_itransformer" else compare_pct(float(row["mse_mean"]), full_mse)
        ablation_rows.append(
            [
                model_name(row["model"]),
                str(int(row["horizon"])),
                fmt(row["mse_mean"], row["mse_std"]),
                fmt(row["mae_mean"], row["mae_std"]),
                mse_delta,
            ]
        )
    add_table(
        doc,
        ["PVG 变体", "预测天数", "MSE 均值±std", "MAE 均值±std", "相对完整 PVG 的 MSE"],
        ablation_rows,
        [1.45, 0.8, 1.45, 1.45, 1.25],
        "表 6 PVG-iTransformer 消融实验结果",
    )

    no_time90_mse = metric(summary, "pvg_no_time", 90, "mse_mean")
    no_time90_mae = metric(summary, "pvg_no_time", 90, "mae_mean")
    no_var90_mse = metric(summary, "pvg_no_variable", 90, "mse_mean")
    no_gate90_mse = metric(summary, "pvg_no_gate", 90, "mse_mean")
    no_time365_mse = metric(summary, "pvg_no_time", 365, "mse_mean")
    no_time365_mae = metric(summary, "pvg_no_time", 365, "mae_mean")
    no_var365_mse = metric(summary, "pvg_no_variable", 365, "mse_mean")
    no_gate365_mse = metric(summary, "pvg_no_gate", 365, "mse_mean")

    paragraph(
        doc,
        f"消融结果显示，变量分支是当前 PVG 结构中贡献最稳定的部分。移除变量分支后，90 天 MSE 从完整 PVG 的 {pvg90_mse:.2f} "
        f"升至 {no_var90_mse:.2f}，即 {compare_pct(no_var90_mse, pvg90_mse)}；365 天 MSE 从 {pvg365_mse:.2f} 升至 "
        f"{no_var365_mse:.2f}，即 {compare_pct(no_var365_mse, pvg365_mse)}。这说明把每个变量的 90 天历史作为 token，"
        "并显式建模目标功率、电力上下文、分表能耗、月度天气和日历变量之间的关系，对本任务是有效的。",
    )
    paragraph(
        doc,
        f"另一方面，移除时间 Patch 分支后模型反而得到最低误差：90 天 MSE 为 {no_time90_mse:.2f}、MAE 为 {no_time90_mae:.2f}，"
        f"相对完整 PVG 的 MSE {compare_pct(no_time90_mse, pvg90_mse)}；365 天 MSE 为 {no_time365_mse:.2f}、MAE 为 "
        f"{no_time365_mae:.2f}，相对完整 PVG 的 MSE {compare_pct(no_time365_mse, pvg365_mse)}。这并不否定时间 Patch 思路，"
        "但说明在当前只有约四年日度样本、天气变量又是月度背景变量的设定下，周级 patch token 可能带来额外参数和局部噪声，"
        f"没有变量 token 视角稳定。移除门控模块后，90 天 MSE 为 {no_gate90_mse:.2f}，与完整 PVG 很接近；365 天 MSE 为 "
        f"{no_gate365_mse:.2f}，相对完整 PVG {compare_pct(no_gate365_mse, pvg365_mse)}，说明可学习门控对长期预测略有帮助，但不是本次性能提升的主要来源。",
    )

    doc.add_heading("3.2 预测曲线对比", level=2)
    paragraph(
        doc,
        "PDF 要求绘制电量预测与 Ground Truth 曲线对比图。以下曲线均取 seed=42 的代表性测试窗口。需要说明的是，曲线图只展示"
        "单个窗口的形态，最终性能比较仍以所有 rolling-origin 测试窗口上的平均 MSE 和 MAE 为准。90 天图可以观察模型对短期上升"
        "趋势和局部波动的跟随情况；365 天图主要观察模型是否能捕捉年度季节轮廓和长期均值水平。",
    )
    figure_idx = 2
    for horizon in [90, 365]:
        for model in ["lstm", "transformer", "pvg_itransformer"]:
            filename = f"{model}_h{horizon}_prediction.png"
            fig_path = WORK_DIR / "figures" / filename
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(6.0))
                add_caption(doc, f"图 {figure_idx} {model_name(model)} {horizon} 天预测曲线")
                figure_idx += 1
    paragraph(
        doc,
        "从代表性曲线可以看到，部分模型在测试起点后的前几十天会高于真实值。这并不说明训练存在错误，而是因为测试集第一个窗口从 "
        "2009 年 7 月 10 日开始，真实用电处于夏季低负荷区间，前 30 天真实均值明显低于训练集均值。模型使用 MSE 训练，倾向于"
        "预测条件均值，因此在低谷阶段容易偏高，在尖峰阶段也可能偏低。这反映了仅依靠历史用电、月度天气和日历变量难以捕捉住户"
        "短期外出或设备使用变化。",
    )

    doc.add_heading("4. 讨论", level=1)
    paragraph(
        doc,
        "本文完成了 PDF 要求的三部分实验。LSTM 作为循环网络基线，结构简单且预测曲线较平滑；标准 Transformer 通过时间 token "
        "和自注意力机制建模较长时间依赖；PVG-iTransformer 则进一步将时间 patch token 与变量 token 结合，并用门控机制自适应"
        "融合两种视角。主实验结果表明，PVG-iTransformer 在 90 天和 365 天任务上均优于 LSTM 与标准 Transformer，因此在当前实验设定下"
        "既具有结构辨识度，也具有较好的实测性能。补充消融实验进一步说明，PVG 的收益主要来自变量 token 分支；完整门控双分支结构并非在所有指标上都优于"
        "只保留变量分支的简化变体。",
    )
    paragraph(
        doc,
        "第三部分模型的主要创新不在于简单增加网络深度，而在于重新定义 Transformer 的输入 token。时间 Patch 分支把连续 90 天"
        "划分为重叠周级片段，适合表达家庭用电中的周周期和近期趋势；变量分支把每个变量的 90 天历史作为 token，适合表达总功率、"
        "电流、电压、分表能耗、月度降水背景和日历变量之间的依赖关系；门控融合则允许模型根据样本自动调整两类信息的权重。"
        "因此，PVG-iTransformer 相比普通时间 token Transformer 更符合本任务的多变量时间序列结构。但从消融结果看，时间 Patch 分支在当前数据规模下并未带来额外性能增益，"
        "这提示第三部分模型的价值应从结构设计和可解释的组件验证两方面评价，而不能只用“完整模型是否绝对最低”来概括。",
    )
    paragraph(
        doc,
        "本实验仍有局限。第一，未获得课程官方 train.csv 和 tes.csv，因此本文按公开 UCI 原始数据自行构造时间划分，结果可能与"
        "官方文件划分不完全一致。第二，天气数据为月度尺度，本文只能将其作为 SURESNES 站点月度气象背景使用，不能刻画每日温度、"
        "降水或湿度的即时影响。第三，NBJBROU 虽在 PDF 中列出，但该站点缺失严重，因此没有纳入主实验；这一选择提升了数据完整性，"
        "但也减少了一个潜在天气变量。第四，365 天预测采用直接多输出方式，输出维度高，训练窗口少，模型容易学习平滑季节轮廓，"
        "对突发峰值或低谷预测不足。",
    )
    paragraph(
        doc,
        "后续可以从四个方向改进：一是使用真实日度天气数据和更近的气象站变量，例如温度、湿度、风速和日降水；二是加入法国节假日、"
        "工作日类型和住户行为代理变量；三是尝试 RevIN、分层预测头或多尺度损失，提高长期预测稳定性；四是进一步做超参数和结构敏感性分析，"
        "例如比较不同 patch_len、stride、门控维度、变量组嵌入方式和输出头设计，以判断时间 Patch 分支在什么条件下可以稳定带来增益。",
    )
    paragraph(
        doc,
        "代码已整理在 /mnt/sdc/zoujunjie/mlearn_power_coursework。正式提交时，应将该目录上传到个人 GitHub 仓库，并把报告中的"
        " GitHub 链接占位替换为真实仓库链接。当前服务器环境没有个人 GitHub 仓库写入权限，因此本文保留可复现代码目录作为说明。"
    )
    small_note(doc, "GitHub 链接占位：https://github.com/<your-id>/mlearn_power_coursework")

    doc.add_heading("参考文献", level=1)
    references = [
        "Dua D, Graff C. UCI Machine Learning Repository: Individual household electric power consumption Data Set. University of California, Irvine.",
        "Meteo-France. Donnees climatologiques de base - mensuelles, MENSQ_92_previous-1950-2024.csv.gz, data.gouv.fr.",
        "Hochreiter S, Schmidhuber J. Long short-term memory. Neural Computation, 1997.",
        "Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need. NeurIPS, 2017.",
        "Nie Y, Nguyen N H, Sinthong P, Kalagnanam J. A Time Series is Worth 64 Words: Long-term Forecasting with Transformers. ICLR, 2023.",
        "Liu Y, Hu T, Zhang H, et al. iTransformer: Inverted Transformers Are Effective for Time Series Forecasting. ICLR, 2024.",
        "课程考核文件：2026 年专硕机器学习课程项目说明。",
        "本文报告文字由 ChatGPT/Codex 辅助整理，实验代码和数值结果由服务器 /mnt/sdc/zoujunjie/mlearn_power_coursework 运行生成。",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(ref)
        east_asia(run)
        run.font.size = Pt(10)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(summary.to_string(index=False))
    print("runs", len(runs))


if __name__ == "__main__":
    main()
